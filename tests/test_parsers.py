"""
tests/test_parsers.py
----------------------
Unit and integration tests for PDFParser and WordParser.
"""
from __future__ import annotations
from pathlib import Path
import pytest
from resume_parser.parsers.pdf_parser import PDFParser
from resume_parser.parsers.word_parser import WordParser


class TestPDFParser:

    def test_extracts_text_from_valid_pdf(self, sample_pdf):
        text = PDFParser().extract_text(str(sample_pdf))
        assert "Jane Doe" in text
        assert "jane.doe@example.com" in text
        assert "Machine Learning" in text

    def test_raises_file_not_found(self, tmp_path):
        with pytest.raises(FileNotFoundError, match="File not found"):
            PDFParser().extract_text(str(tmp_path / "nonexistent.pdf"))

    def test_raises_value_error_for_wrong_extension(self, tmp_path):
        f = tmp_path / "resume.txt"
        f.write_text("irrelevant")
        with pytest.raises(ValueError, match="PDFParser expects a '.pdf' file"):
            PDFParser().extract_text(str(f))

    def test_raises_value_error_for_corrupt_pdf(self, tmp_path):
        f = tmp_path / "bad.pdf"
        f.write_bytes(b"NOT_A_PDF")
        with pytest.raises(ValueError, match="PDFParser could not parse"):
            PDFParser().extract_text(str(f))

    def test_strip_whitespace_false_still_returns_text(self, sample_pdf):
        text = PDFParser(strip_whitespace=False).extract_text(str(sample_pdf))
        assert len(text) > 0

    def test_returns_string(self, sample_pdf):
        assert isinstance(PDFParser().extract_text(str(sample_pdf)), str)


class TestWordParser:

    def test_extracts_text_from_valid_docx(self, sample_docx):
        text = WordParser().extract_text(str(sample_docx))
        assert "Jane Doe" in text
        assert "jane.doe@example.com" in text
        assert "PyTorch" in text

    def test_raises_file_not_found(self, tmp_path):
        with pytest.raises(FileNotFoundError, match="File not found"):
            WordParser().extract_text(str(tmp_path / "ghost.docx"))

    def test_raises_value_error_for_wrong_extension(self, tmp_path):
        f = tmp_path / "doc.pdf"
        f.write_bytes(b"%PDF-1.4")
        with pytest.raises(ValueError, match="WordParser expects a '.docx' file"):
            WordParser().extract_text(str(f))

    def test_raises_value_error_for_corrupt_docx(self, tmp_path):
        f = tmp_path / "broken.docx"
        f.write_bytes(b"DEFINITELY_NOT_A_ZIP")
        with pytest.raises(ValueError, match="WordParser could not open"):
            WordParser().extract_text(str(f))

    def test_include_tables_extracts_cell_text(self, tmp_path):
        import docx as docx_lib
        path = tmp_path / "table.docx"
        doc = docx_lib.Document()
        doc.add_paragraph("Alice Smith")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Python"
        table.rows[0].cells[1].text = "AWS"
        doc.save(str(path))
        text = WordParser(include_tables=True).extract_text(str(path))
        assert "Python" in text
        assert "AWS" in text

    def test_include_tables_false_omits_cells(self, tmp_path):
        import docx as docx_lib
        path = tmp_path / "table.docx"
        doc = docx_lib.Document()
        doc.add_paragraph("Alice Smith")
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "OnlyInTable"
        doc.save(str(path))
        text = WordParser(include_tables=False).extract_text(str(path))
        assert "Alice Smith" in text
        assert "OnlyInTable" not in text

    def test_returns_string(self, sample_docx):
        assert isinstance(WordParser().extract_text(str(sample_docx)), str)