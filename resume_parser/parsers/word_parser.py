"""
resume_parser/parsers/word_parser.py
--------------------------------------
Concrete FileParser for Word (.docx) documents using python-docx.
"""
from __future__ import annotations
import logging
from pathlib import Path
import docx
from resume_parser.base import FileParser

logger = logging.getLogger(__name__)


class WordParser(FileParser):
    """
    Extract plain text from a Word (.docx) resume using python-docx.

    Collects text from body paragraphs and optionally table cells.

    Parameters
    ----------
    include_tables:
        When True (default), text from table cells is appended after paragraphs.
    """

    def __init__(self, *, include_tables: bool = True) -> None:
        self._include_tables = include_tables

    def extract_text(self, file_path: str) -> str:
        """
        Read file_path (must be .docx) and return full text content.

        Raises
        ------
        FileNotFoundError  – if file_path does not exist.
        ValueError         – if not a .docx or python-docx cannot open it.
        """
        path = Path(file_path)
        self._validate_path(path)
        logger.info("WordParser: opening '%s'", path)
        try:
            document = docx.Document(str(path))
        except Exception as exc:
            raise ValueError(f"WordParser could not open '{path}': {exc}") from exc

        lines: list[str] = [p.text for p in document.paragraphs]
        if self._include_tables:
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        lines.append(cell.text)

        full_text = "\n".join(line for line in lines if line.strip())
        logger.info("WordParser: extracted %d chars from '%s'.", len(full_text), path.name)
        return full_text

    @staticmethod
    def _validate_path(path: Path) -> None:
        if not path.exists():
            raise FileNotFoundError(f"File not found: '{path}'")
        if path.suffix.lower() != ".docx":
            raise ValueError(f"WordParser expects a '.docx' file, got '{path.suffix}'")